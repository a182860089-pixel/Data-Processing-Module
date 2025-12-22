"""
Word格式化器
将Markdown内容转换为Word文档，保持格式一致
"""
import re
import logging
from pathlib import Path
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

import markdown
from bs4 import BeautifulSoup, NavigableString

logger = logging.getLogger(__name__)


class WordFormatter:
    """
    Markdown转Word格式化器
    
    支持的格式：
    - 标题层级 (H1-H6)
    - 段落和换行
    - 粗体/斜体
    - 有序/无序列表
    - 表格（带边框和样式）
    - 代码块
    - 引用块
    """
    
    def __init__(self):
        """初始化格式化器"""
        self.default_font_name = "Microsoft YaHei"  # 默认字体：微软雅黑
        self.default_font_size = Pt(11)
        self.code_font_name = "Consolas"
        self.code_font_size = Pt(10)
    
    def markdown_to_docx(
        self, 
        md_content: str, 
        output_path: str,
        title: Optional[str] = None
    ) -> str:
        """
        将Markdown或HTML内容转换为Word文档
        
        Args:
            md_content: Markdown或HTML内容
            output_path: 输出文件路径
            title: 文档标题（可选）
            
        Returns:
            str: 输出文件路径
        """
        # 创建Word文档
        doc = Document()
        
        # 设置默认样式
        self._set_default_styles(doc)
        
        # 添加标题（如果提供）
        if title:
            doc.add_heading(title, level=0)
        
        # 检测内容是否已经是HTML（OCR可能直接返回HTML表格）
        content = md_content.strip()
        is_html = bool(re.search(r'<(table|tr|td|th|div|p|h[1-6]|ul|ol|li)\b', content, re.IGNORECASE))
        
        if is_html:
            # 内容已经是HTML，直接解析
            html = content
            logger.info("Detected HTML content, skipping markdown conversion")
            # 解析HTML
            soup = BeautifulSoup(html, 'html.parser')
            # 处理HTML元素
            self._process_elements(doc, soup)
        else:
            # 检测是否包含 Markdown 表格，使用自定义解析
            if self._contains_markdown_table(content):
                logger.info("Detected Markdown table, using custom parser")
                self._process_markdown_content(doc, content)
            else:
                # 将Markdown转换为HTML
                html = markdown.markdown(
                    content, 
                    extensions=['tables', 'fenced_code', 'nl2br']
                )
                # 解析HTML
                soup = BeautifulSoup(html, 'html.parser')
                # 处理HTML元素
                self._process_elements(doc, soup)
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"Word document saved: {output_path}")
        
        return output_path
    
    def _contains_markdown_table(self, content: str) -> bool:
        """
        检测内容是否包含 Markdown 表格
        
        Args:
            content: Markdown 内容
            
        Returns:
            bool: 是否包含表格
        """
        lines = content.split('\n')
        for line in lines:
            # Markdown 表格分隔行的特征：包含 | 和 - 的组合
            if re.match(r'^\s*\|?[\s\-:|]+\|[\s\-:|]+\|?\s*$', line):
                return True
        return False
    
    def _process_markdown_content(self, doc: Document, content: str):
        """
        处理 Markdown 内容，特别处理表格
        
        Args:
            doc: Word 文档对象
            content: Markdown 内容
        """
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 检测表格开始（包含 | 的行）
            if '|' in line and i + 1 < len(lines):
                # 查找表格分隔行
                next_line = lines[i + 1] if i + 1 < len(lines) else ''
                if re.match(r'^\s*\|?[\s\-:|]+\|[\s\-:|]+\|?\s*$', next_line):
                    # 找到表格，收集所有表格行
                    table_lines = [line]
                    i += 1
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i])
                        i += 1
                    # 处理表格
                    self._process_markdown_table(doc, table_lines)
                    continue
            
            # 处理标题
            if line.startswith('#'):
                match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if match:
                    level = len(match.group(1))
                    text = match.group(2).strip()
                    heading = doc.add_heading(text, level=level)
                    self._set_chinese_font(heading)
                    i += 1
                    continue
            
            # 处理普通段落
            stripped = line.strip()
            if stripped:
                para = doc.add_paragraph()
                # 处理粗体和斜体
                self._process_markdown_inline(para, stripped)
            
            i += 1
    
    def _process_markdown_inline(self, paragraph, text: str):
        """
        处理 Markdown 行内格式（粗体、斜体等）
        
        Args:
            paragraph: Word 段落对象
            text: 文本内容
        """
        # 简化处理：直接添加文本
        # 如果需要更复杂的格式处理，可以在这里扩展
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),      # **粗体**
            (r'__(.+?)__', 'bold'),           # __粗体__
            (r'\*(.+?)\*', 'italic'),         # *斜体*
            (r'_(.+?)_', 'italic'),           # _斜体_
        ]
        
        # 简单实现：直接添加文本
        run = paragraph.add_run(text)
        self._set_run_font(run)
    
    def _process_markdown_table(self, doc: Document, table_lines: List[str]):
        """
        处理 Markdown 表格
        
        Args:
            doc: Word 文档对象
            table_lines: 表格行列表
        """
        if len(table_lines) < 2:
            return
        
        # 解析表格数据
        rows_data = []
        header_row = None
        separator_idx = -1
        
        for idx, line in enumerate(table_lines):
            # 跳过空行
            if not line.strip():
                continue
            
            # 检测分隔行
            if re.match(r'^\s*\|?[\s\-:|]+\|[\s\-:|]+\|?\s*$', line):
                separator_idx = idx
                continue
            
            # 解析单元格
            cells = self._parse_table_row(line)
            if cells:
                if separator_idx == -1:
                    # 分隔行之前是表头
                    header_row = cells
                else:
                    rows_data.append(cells)
        
        # 如果有表头，将其作为第一行
        if header_row:
            rows_data.insert(0, header_row)
        
        if not rows_data:
            return
        
        # 计算最大列数
        col_count = max(len(row) for row in rows_data)
        row_count = len(rows_data)
        
        logger.info(f"Creating table with {row_count} rows and {col_count} columns")
        
        # 创建表格
        table = doc.add_table(rows=row_count, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格边框
        self._set_table_borders(table)
        
        # 填充数据
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < col_count:
                    cell = table.cell(i, j)
                    cell.text = cell_text.strip()
                    
                    # 设置样式
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            self._set_run_font(run)
                            # 表头行加粗
                            if i == 0 and header_row:
                                run.bold = True
        
        # 添加表格后的空行
        doc.add_paragraph()
    
    def _parse_table_row(self, line: str) -> List[str]:
        """
        解析表格行
        
        Args:
            line: 表格行文本
            
        Returns:
            List[str]: 单元格列表
        """
        # 移除首尾的 |
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        
        # 按 | 分割
        cells = line.split('|')
        return [cell.strip() for cell in cells]
    
    def _set_default_styles(self, doc: Document):
        """
        设置文档默认样式
        
        Args:
            doc: Word文档对象
        """
        # 设置正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = self.default_font_name
        font.size = self.default_font_size
        
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font_name)
    
    def _process_elements(self, doc: Document, soup: BeautifulSoup):
        """
        处理HTML元素并转换为Word内容
        
        Args:
            doc: Word文档对象
            soup: BeautifulSoup对象
        """
        for element in soup.children:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    doc.add_paragraph(text)
            elif element.name:
                self._process_element(doc, element)
    
    def _process_element(self, doc: Document, element):
        """
        处理单个HTML元素
        
        Args:
            doc: Word文档对象
            element: HTML元素
        """
        tag_name = element.name
        
        # 处理标题
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            heading = doc.add_heading(element.get_text(), level=level)
            self._set_chinese_font(heading)
        
        # 处理段落
        elif tag_name == 'p':
            para = doc.add_paragraph()
            self._process_inline_elements(para, element)
        
        # 处理无序列表
        elif tag_name == 'ul':
            self._process_list(doc, element, ordered=False)
        
        # 处理有序列表
        elif tag_name == 'ol':
            self._process_list(doc, element, ordered=True)
        
        # 处理表格
        elif tag_name == 'table':
            self._process_table(doc, element)
        
        # 处理代码块
        elif tag_name == 'pre':
            self._process_code_block(doc, element)
        
        # 处理引用块
        elif tag_name == 'blockquote':
            self._process_blockquote(doc, element)
        
        # 处理分隔线
        elif tag_name == 'hr':
            doc.add_paragraph('─' * 50)
        
        # 处理其他块级元素
        elif tag_name in ['div', 'section', 'article']:
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    self._process_element(doc, child)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        doc.add_paragraph(text)
    
    def _process_inline_elements(self, paragraph, element):
        """
        处理行内元素（粗体、斜体等）
        
        Args:
            paragraph: Word段落对象
            element: HTML元素
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    self._set_run_font(run)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
                self._set_run_font(run)
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
                self._set_run_font(run)
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = self.code_font_name
                run.font.size = self.code_font_size
                # 设置灰色背景效果（通过高亮）
                run.font.highlight_color = 15  # 浅灰色
            elif child.name == 'a':
                # 处理链接 - 显示为蓝色下划线文本
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = True
                self._set_run_font(run)
            elif child.name == 'br':
                paragraph.add_run('\n')
            else:
                # 递归处理其他行内元素
                self._process_inline_elements(paragraph, child)
    
    def _process_list(self, doc: Document, element, ordered: bool = False):
        """
        处理列表
        
        Args:
            doc: Word文档对象
            element: 列表HTML元素
            ordered: 是否为有序列表
        """
        style = 'List Number' if ordered else 'List Bullet'
        
        for i, li in enumerate(element.find_all('li', recursive=False)):
            para = doc.add_paragraph(style=style)
            
            # 处理列表项内容
            for child in li.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        run = para.add_run(text)
                        self._set_run_font(run)
                elif child.name in ['strong', 'b']:
                    run = para.add_run(child.get_text())
                    run.bold = True
                    self._set_run_font(run)
                elif child.name in ['em', 'i']:
                    run = para.add_run(child.get_text())
                    run.italic = True
                    self._set_run_font(run)
                elif child.name == 'code':
                    run = para.add_run(child.get_text())
                    run.font.name = self.code_font_name
                    run.font.size = self.code_font_size
                elif child.name in ['ul', 'ol']:
                    # 处理嵌套列表
                    self._process_list(doc, child, ordered=(child.name == 'ol'))
    
    def _process_table(self, doc: Document, element):
        """
        处理表格（支持colspan和rowspan合并单元格）
        
        Args:
            doc: Word文档对象
            element: 表格HTML元素
        """
        rows = element.find_all('tr')
        if not rows:
            return
        
        # 第一遍：分析表格结构，计算实际列数
        col_count = 0
        row_col_counts = []
        for row in rows:
            row_cols = 0
            for cell in row.find_all(['td', 'th']):
                colspan = int(cell.get('colspan', 1))
                row_cols += colspan
            row_col_counts.append(row_cols)
            col_count = max(col_count, row_cols)
        
        if col_count == 0:
            return
        
        logger.info(f"Processing HTML table: {len(rows)} rows, {col_count} cols")
        
        # 如果列数较多，设置页面为横向
        if col_count > 8:
            self._set_landscape_orientation(doc)
        
        # 计算实际行数
        row_count = len(rows)
        
        # 创建表格
        table = doc.add_table(rows=row_count, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格边框
        self._set_table_borders(table)
        
        # 设置表格和列宽
        self._set_table_width(table, col_count)
        
        # 跟踪单元格占用情况（用于处理rowspan）
        occupied = [[False] * col_count for _ in range(row_count)]
        
        # 根据列数确定字体大小
        if col_count > 12:
            table_font_size = Pt(8)
        elif col_count > 8:
            table_font_size = Pt(9)
        else:
            table_font_size = Pt(10)
        
        # 填充表格内容
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            col_idx = 0
            cell_idx = 0
            
            while col_idx < col_count and cell_idx < len(cells):
                # 跳过已被占用的单元格
                if occupied[i][col_idx]:
                    col_idx += 1
                    continue
                
                cell = cells[cell_idx]
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))
                
                # 如果当前行列数不足，且这是最后几个单元格，需要扩展colspan来填满
                remaining_cols = col_count - col_idx
                
                # 确保不超出边界
                colspan = min(colspan, remaining_cols)
                rowspan = min(rowspan, row_count - i)
                
                # 如果这是当前行的最后一个单元格，且还有剩余列，扩展colspan
                if cell_idx == len(cells) - 1 and row_col_counts[i] < col_count:
                    # 计算需要额外合并的列数
                    extra_cols = col_count - row_col_counts[i]
                    colspan = min(colspan + extra_cols, remaining_cols)
                
                # 获取单元格内容（处理br标签）
                cell_text = self._get_cell_text(cell)
                
                # 处理单元格合并
                if colspan > 1 or rowspan > 1:
                    try:
                        start_cell = table.cell(i, col_idx)
                        end_cell = table.cell(i + rowspan - 1, col_idx + colspan - 1)
                        start_cell.merge(end_cell)
                    except Exception as e:
                        logger.warning(f"Failed to merge cells at row {i}, col {col_idx}: {e}")
                
                # 填充内容
                table_cell = table.cell(i, col_idx)
                table_cell.text = cell_text
                
                # 设置单元格样式
                for paragraph in table_cell.paragraphs:
                    # 减少段落间距
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    for run in paragraph.runs:
                        run.font.name = self.default_font_name
                        run.font.size = table_font_size
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font_name)
                        # 表头加粗
                        if cell.name == 'th':
                            run.bold = True
                    # 居中对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 标记占用的单元格
                for r in range(rowspan):
                    for c in range(colspan):
                        if i + r < row_count and col_idx + c < col_count:
                            occupied[i + r][col_idx + c] = True
                
                col_idx += colspan
                cell_idx += 1
        
        # 添加表格后的空行
        doc.add_paragraph()
    
    def _set_landscape_orientation(self, doc: Document):
        """
        设置页面为横向
        
        Args:
            doc: Word文档对象
        """
        try:
            section = doc.sections[-1]
            # 获取当前宽高
            new_width = section.page_height
            new_height = section.page_width
            # 设置横向
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            # 减小页边距
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            logger.info("Set page to landscape orientation")
        except Exception as e:
            logger.warning(f"Failed to set landscape orientation: {e}")
    
    def _set_table_width(self, table, col_count: int):
        """
        设置表格和列宽
        
        Args:
            table: Word表格对象
            col_count: 列数
        """
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # 设置表格宽度为100%
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')  # 百分比
            tblPr.append(tblW)
            
            # 设置表格布局为固定
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # 计算每列宽度（平均分配）
            # A4横向可用宽度约为27cm（29.7 - 2*1.35边距）
            # 转换为 twips (1 cm = 567 twips)
            total_width = 27 * 567  # twips
            col_width = int(total_width / col_count)
            
            # 设置每列宽度
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(col_width))
                    tcW.set(qn('w:type'), 'dxa')  # twips
                    tcPr.append(tcW)
                    
        except Exception as e:
            logger.warning(f"Failed to set table width: {e}")
    
    def _get_cell_text(self, cell) -> str:
        """
        获取单元格文本内容，处理br标签
        
        Args:
            cell: HTML单元格元素
            
        Returns:
            str: 单元格文本
        """
        # 将br标签替换为换行符
        for br in cell.find_all('br'):
            br.replace_with('\n')
        return cell.get_text().strip()
    
    def _set_table_borders(self, table):
        """
        设置表格边框样式
        
        Args:
            table: Word表格对象
        """
        tbl = table._tbl
        
        # 获取或创建 tblPr 元素
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # 移除已有的边框设置（如果有）
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        
        # 创建边框元素
        tblBorders = OxmlElement('w:tblBorders')
        
        # 设置各个边框
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')  # 单线边框
            border.set(qn('w:sz'), '4')        # 边框粗细
            border.set(qn('w:space'), '0')     # 边框间距
            border.set(qn('w:color'), '000000')  # 黑色边框
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _process_code_block(self, doc: Document, element):
        """
        处理代码块
        
        Args:
            doc: Word文档对象
            element: 代码块HTML元素
        """
        code_element = element.find('code')
        if code_element:
            code_text = code_element.get_text()
        else:
            code_text = element.get_text()
        
        # 创建代码块段落
        para = doc.add_paragraph()
        
        # 设置代码块样式
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # 添加代码文本
        run = para.add_run(code_text)
        run.font.name = self.code_font_name
        run.font.size = self.code_font_size
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.code_font_name)
    
    def _process_blockquote(self, doc: Document, element):
        """
        处理引用块
        
        Args:
            doc: Word文档对象
            element: 引用块HTML元素
        """
        para = doc.add_paragraph()
        
        # 设置引用块样式
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # 添加引用文本
        text = element.get_text().strip()
        run = para.add_run(f"「{text}」")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
        self._set_run_font(run)
    
    def _set_run_font(self, run):
        """
        设置run的字体
        
        Args:
            run: Word run对象
        """
        run.font.name = self.default_font_name
        run.font.size = self.default_font_size
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font_name)
    
    def _set_chinese_font(self, paragraph):
        """
        设置段落的中文字体
        
        Args:
            paragraph: Word段落对象
        """
        for run in paragraph.runs:
            run.font.name = self.default_font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font_name)


class ImageToWordConverter:
    """
    图片转Word转换器
    
    流程：图片 -> OCR识别 -> Markdown -> Word文档
    """
    
    def __init__(self):
        """初始化转换器"""
        self.word_formatter = WordFormatter()
    
    async def convert(
        self,
        image_path: str,
        output_path: str,
        ocr_client,
        title: Optional[str] = None
    ) -> str:
        """
        将图片转换为Word文档
        
        Args:
            image_path: 图片文件路径
            output_path: 输出Word文件路径
            ocr_client: OCR客户端实例
            title: 文档标题（可选）
            
        Returns:
            str: 输出文件路径
        """
        import base64
        
        # 1. 读取图片并转为Base64
        with open(image_path, 'rb') as f:
            image_data = f.read()
        image_base64 = base64.b64encode(image_data).decode('utf-8')
        
        # 2. 调用OCR获取Markdown
        markdown_content = await ocr_client.ocr_image(image_base64)
        
        logger.info(f"OCR completed, markdown length: {len(markdown_content)}")
        
        # 3. 将Markdown转换为Word
        result_path = self.word_formatter.markdown_to_docx(
            markdown_content,
            output_path,
            title=title
        )
        
        return result_path
    
    async def convert_from_base64(
        self,
        image_base64: str,
        output_path: str,
        ocr_client,
        title: Optional[str] = None
    ) -> Tuple[str, str]:
        """
        从Base64图片转换为Word文档
        
        Args:
            image_base64: Base64编码的图片
            output_path: 输出Word文件路径
            ocr_client: OCR客户端实例
            title: 文档标题（可选）
            
        Returns:
            Tuple[str, str]: (输出文件路径, Markdown内容)
        """
        # 1. 调用OCR获取Markdown
        markdown_content = await ocr_client.ocr_image(image_base64)
        
        logger.info(f"OCR completed, markdown length: {len(markdown_content)}")
        
        # 2. 将Markdown转换为Word
        result_path = self.word_formatter.markdown_to_docx(
            markdown_content,
            output_path,
            title=title
        )
        
        return result_path, markdown_content
